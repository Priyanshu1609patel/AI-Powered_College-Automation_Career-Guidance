"""
Resume Parser — robust section-based + full-text fallback extraction.
Works for any PDF/DOCX resume regardless of formatting style.
"""
import re
import fitz          # PyMuPDF
from docx import Document

# ── spaCy (optional, gracefully skipped) ─────────────────────────────────────
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except Exception:
    nlp = None
    SPACY_AVAILABLE = False

# ── Comprehensive skill vocabulary ───────────────────────────────────────────
SKILLS_VOCAB = [
    # Languages
    'python','java','c++','c#','c','javascript','typescript','ruby','php',
    'swift','kotlin','go','rust','scala','r','matlab','perl','dart',
    'bash','shell scripting','powershell','groovy','assembly',
    # Frontend
    'html','css','react','angular','vue','vue.js','react.js','next.js',
    'nuxt.js','bootstrap','tailwind','jquery','redux','webpack','sass',
    'less','svelte','flutter',
    # Backend
    'node.js','nodejs','express','flask','django','fastapi','spring',
    'spring boot','laravel','rails','asp.net','graphql','rest api',
    'restful api','microservices','api development','servlet','hibernate',
    # Databases
    'sql','mysql','postgresql','mongodb','redis','sqlite','oracle',
    'nosql','firebase','dynamodb','cassandra','elasticsearch','mariadb',
    'ms sql server','supabase','jdbc',
    # Cloud / DevOps
    'aws','azure','gcp','google cloud','docker','kubernetes','jenkins',
    'ci/cd','terraform','ansible','linux','unix','git','github','gitlab',
    'bitbucket','nginx','apache','heroku','vercel','netlify',
    # Data / AI / ML
    'machine learning','deep learning','data analysis','data science',
    'tensorflow','pytorch','keras','pandas','numpy','scikit-learn',
    'natural language processing','nlp','computer vision','opencv',
    'tableau','power bi','hadoop','spark','etl','data visualization',
    'statistics','data mining','big data','feature engineering',
    'matplotlib','seaborn','scipy','xgboost','lstm','bert',
    'neural network','regression','classification','clustering',
    # Security
    'cybersecurity','ethical hacking','penetration testing',
    'network security','cryptography','kali linux',
    # Tools
    'jira','figma','postman','vs code','excel','photoshop','linux',
    'networking','agile','scrum','kanban','confluence','trello',
    'microsoft office','ms office','autocad','solidworks','jupyter',
    # Soft skills
    'leadership','communication','teamwork','problem solving',
    'project management','time management','critical thinking',
    'collaboration','presentation','negotiation','mentoring',
    'public speaking','analytical thinking',
]

# ── Section-keyword sets (use "contains" matching, not exact) ─────────────────
_SKILL_KWS = ['skill', 'competenc', 'expertise', 'technolog', 'tool',
               'proficien', 'programming language', 'technical']
_EDU_KWS   = ['education', 'academic', 'qualification', 'degree', 'study',
               'studies', 'schooling', 'course']
_EXP_KWS   = ['experience', 'employment', 'work history', 'career',
               'internship', 'project', 'professional', 'position',
               'responsibility', 'responsibilit', 'achievement']
_OTHER_KWS = ['certification', 'award', 'interest', 'hobby', 'language',
              'reference', 'declaration', 'objective', 'summary', 'profile',
              'contact', 'personal', 'extracurricular', 'volunteer']

# ── Degree patterns ───────────────────────────────────────────────────────────
_DEGREE_RE = re.compile(
    r'\b(b\.?\s*tech|b\.?\s*e\.?|b\.?\s*sc\.?|b\.?\s*com\.?|b\.?\s*a\.?'
    r'|bca|mca|m\.?\s*tech|m\.?\s*e\.?|m\.?\s*sc\.?|m\.?\s*com\.?'
    r'|m\.?\s*a\.?|mba|bba|phd|ph\.?\s*d\.?|bachelor|master|diploma'
    r'|10th|12th|ssc|hsc|class\s*x\b|class\s*xii|secondary|senior secondary'
    r'|undergraduate|postgraduate|associate)',
    re.IGNORECASE
)
_INST_RE = re.compile(
    r'\b(university|college|institute|school|iit|nit|bits|polytechnic'
    r'|academy|mit|faculty)\b',
    re.IGNORECASE
)

# ─────────────────────────────────────────────────────────────────────────────
#  Text extraction
# ─────────────────────────────────────────────────────────────────────────────
def extract_text(file_path):
    if file_path.lower().endswith('.pdf'):
        return _pdf_text(file_path)
    if file_path.lower().endswith('.docx'):
        return _docx_text(file_path)
    return ''

def _pdf_text(path):
    text = ''
    try:
        doc = fitz.open(path)
        for page in doc:
            text += page.get_text('text') + '\n'
    except Exception:
        pass
    return text

def _docx_text(path):
    try:
        doc = Document(path)
        lines = [p.text for p in doc.paragraphs]
        # also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    lines.append(cell.text)
        return '\n'.join(lines)
    except Exception:
        return ''

# ─────────────────────────────────────────────────────────────────────────────
#  Section splitter  — uses "contains" matching so "TECHNICAL SKILLS & TOOLS"
#  is correctly identified as a skills header.
# ─────────────────────────────────────────────────────────────────────────────
def _looks_like_header(line):
    """Heuristic: a header is short, has no sentence punctuation, may be ALL CAPS."""
    stripped = line.strip()
    if not stripped or len(stripped) > 60:
        return False
    # all-caps or title-case short line
    words = stripped.split()
    if len(words) > 6:
        return False
    return True

def _classify_header(lower):
    """Return section tag if the line looks like a known section header."""
    if any(kw in lower for kw in _SKILL_KWS):
        return 'skills'
    if any(kw in lower for kw in _EDU_KWS):
        return 'education'
    if any(kw in lower for kw in _EXP_KWS):
        return 'experience'
    if any(kw in lower for kw in _OTHER_KWS):
        return 'other'
    return None

def split_sections(text):
    """
    Returns {'skills': [...lines], 'education': [...], 'experience': [...], 'general': [...]}
    Lines before the first recognised header go into 'general'.
    """
    sections = {'skills': [], 'education': [], 'experience': [], 'general': []}
    current = 'general'

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        lower = re.sub(r'[:\-_=•|]', ' ', line.lower()).strip()
        if _looks_like_header(line):
            tag = _classify_header(lower)
            if tag and tag != 'other':
                current = tag
                continue
            elif tag == 'other':
                current = 'general'
                continue
        sections[current].append(line)

    return sections

# ─────────────────────────────────────────────────────────────────────────────
#  Field extractors
# ─────────────────────────────────────────────────────────────────────────────
def _extract_skills(text, sections):
    found = set()
    full_lower = text.lower()

    # 1. Full-text keyword match (catches skills mentioned anywhere)
    for skill in SKILLS_VOCAB:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, full_lower):
            found.add(skill.title())

    # 2. Tokenise skills-section lines (catches "Python | Java | React" style)
    for line in sections.get('skills', []):
        for token in re.split(r'[,|/•·\-–—\t]+', line):
            t = token.strip().lower()
            if 1 < len(t) <= 35:
                for skill in SKILLS_VOCAB:
                    if skill == t or (len(skill) > 4 and skill in t):
                        found.add(skill.title())

    return ', '.join(sorted(found))


def _extract_education(text, sections):
    lines = []

    # Strategy 1 — lines inside detected education section
    for l in sections.get('education', []):
        if len(l) > 3:
            lines.append(l)

    # Strategy 2 — full-text fallback: look for degree or institution lines
    if not lines:
        for l in text.splitlines():
            l = l.strip()
            if not l:
                continue
            if _DEGREE_RE.search(l) or _INST_RE.search(l):
                lines.append(l)

    # De-duplicate, preserve order
    seen, result = set(), []
    for l in lines:
        if l not in seen:
            seen.add(l)
            result.append(l)

    return ' | '.join(result[:8])


def _extract_experience(text, sections):
    lines = []

    # Strategy 1 — lines inside detected experience / project sections
    for l in sections.get('experience', []):
        if len(l.strip()) > 5:
            lines.append(l.strip())

    # Strategy 2 — full-text fallback: year-range lines or action-verb sentences
    if not lines:
        year_re  = re.compile(r'\b(19|20)\d{2}\b')
        verbs_re = re.compile(
            r'\b(developed|built|designed|managed|led|created|implemented'
            r'|analyzed|improved|launched|deployed|worked|interned|contributed'
            r'|maintained|automated|optimized|delivered|engineered|resolved)\b',
            re.IGNORECASE
        )
        for l in text.splitlines():
            l = l.strip()
            if len(l) < 8:
                continue
            if verbs_re.search(l) or year_re.search(l):
                lines.append(l)

    seen, result = set(), []
    for l in lines:
        if l not in seen:
            seen.add(l)
            result.append(l)

    return ' | '.join(result[:10])


# ─────────────────────────────────────────────────────────────────────────────
#  Main entry point
# ─────────────────────────────────────────────────────────────────────────────
def extract_resume_data(text):
    if not text or not text.strip():
        return {'skills': '', 'education': '', 'experience': '', 'raw_text': ''}

    sections  = split_sections(text)
    skills    = _extract_skills(text, sections)
    education = _extract_education(text, sections)
    experience = _extract_experience(text, sections)

    return {
        'skills':     skills,
        'education':  education,
        'experience': experience,
        'raw_text':   text,
    }
